"""DOCX 导入/导出工具模块。

编辑器标记 ←→ DOCX 格式映射：
  \\f（分页符）       ←→  Word 分页符
  　　 （全角空格缩进） ←→  段落首行缩进
"""

import io
import re
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 导出手稿样式常量 ────────────────────────────────────────────

FONT_BODY = '宋体'
FONT_TITLE = '黑体'
SIZE_BODY = Pt(12)       # 小四
SIZE_TITLE = Pt(16)      # 三号
LINE_SPACING = 1.5
INDENT_CHARS = 2         # 段首缩进字符数


def _set_run_font(run, name=FONT_BODY, size=SIZE_BODY, bold=False):
    """设置 run 的字体"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = size
    run.bold = bold


def _set_paragraph_spacing(paragraph, line_spacing=LINE_SPACING):
    """设置段落行距"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing


def _add_page_break(paragraph):
    """在段落前插入分页符"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run('')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)


def _estimate_indent_width(paragraph):
    """估算两字符缩进的实际宽度（基于正文字号）"""
    # 首行缩进 = 字号磅值 × 2 字符
    return SIZE_BODY * INDENT_CHARS


def _has_fullwidth_indent(text):
    """检测段落是否以全角空格缩进开头"""
    return text.startswith('　　') or text.startswith('　')


def _strip_indent(text):
    """去掉段落开头的全角空格缩进"""
    return text.lstrip('　')


# ── 导出 ────────────────────────────────────────────────────────

def build_manuscript_docx(title, content, chapter_title=None):
    """从编辑器内容构建一份标准手稿 DOCX。

    参数：
      title: 文档标题（书名）
      content: 编辑器原始文本（含 \\f 分页符和全角空格缩进）
      chapter_title: 可选，章节标题（全书合并时可为 None）

    返回：python-docx Document 对象
    """
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ── 默认段落样式 ──
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = SIZE_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # ── 书名标题 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(title_para)
    title_run = title_para.add_run(title)
    _set_run_font(title_run, FONT_TITLE, SIZE_TITLE, bold=True)

    # ── 章节标题（如有） ──
    if chapter_title:
        ch_para = doc.add_paragraph()
        ch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(ch_para)
        ch_run = ch_para.add_run(chapter_title)
        _set_run_font(ch_run, FONT_TITLE, Pt(14), bold=True)

    # ── 正文段落 ──
    segments = content.split('\f')
    for i, seg in enumerate(segments):
        if i > 0:
            # 分页符：在当前段落末尾加 page break
            if doc.paragraphs:
                _add_page_break(doc.paragraphs[-1])

        lines = seg.strip().split('\n')
        for line in lines:
            if not line.strip():
                continue
            para = doc.add_paragraph()
            _set_paragraph_spacing(para)

            stripped = _strip_indent(line)
            has_indent = _has_fullwidth_indent(line)

            run = para.add_run(stripped)
            _set_run_font(run, FONT_BODY, SIZE_BODY)

            if has_indent:
                pf = para.paragraph_format
                pf.first_line_indent = _estimate_indent_width(para)

    return doc


def build_chapter_docx(chapter):
    """从 Chapter 模型构建单章 DOCX"""
    title = chapter.book.title if chapter.book else ''
    return build_manuscript_docx(title, chapter.content or '', chapter.title)


def build_book_docx(book):
    """从 Book 模型构建全书合并 DOCX"""
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ── 默认段落样式 ──
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = SIZE_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # ── 书名标题 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(title_para)
    title_run = title_para.add_run(book.title)
    _set_run_font(title_run, FONT_TITLE, SIZE_TITLE, bold=True)

    # ── 逐章 ──
    for idx, chapter in enumerate(book.chapters):
        if idx > 0:
            # 章间用分页符
            if doc.paragraphs:
                _add_page_break(doc.paragraphs[-1])

        # 章节标题
        ch_para = doc.add_paragraph()
        ch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(ch_para)
        ch_run = ch_para.add_run(chapter.title)
        _set_run_font(ch_run, FONT_TITLE, Pt(14), bold=True)

        # 正文
        content = chapter.content or ''
        segments = content.split('\f')
        for j, seg in enumerate(segments):
            if j > 0:
                if doc.paragraphs:
                    _add_page_break(doc.paragraphs[-1])

            lines = seg.strip().split('\n')
            for line in lines:
                if not line.strip():
                    continue
                para = doc.add_paragraph()
                _set_paragraph_spacing(para)

                stripped = _strip_indent(line)
                has_indent = _has_fullwidth_indent(line)

                run = para.add_run(stripped)
                _set_run_font(run, FONT_BODY, SIZE_BODY)

                if has_indent:
                    pf = para.paragraph_format
                    pf.first_line_indent = _estimate_indent_width(para)

    return doc


# ── 导入 ────────────────────────────────────────────────────────

def parse_docx(filepath_or_stream):
    """解析 DOCX 文件，提取文本和结构。

    返回：
      {
        'title': str | None,       # 文档中检测到的书名/标题
        'chapters': [               # 按 Heading 拆分的章节列表
          {'title': str, 'content': str},
        ]
      }

    映射规则：
      - Heading 1 段落 → 章节分界，作为章节标题
      - Word 分页符 → \\f（同章内分页）
      - 段首缩进 → 　　 （两个全角空格）
      - 普通段落 → 保留文本
    """
    doc = Document(filepath_or_stream)

    chapters = []
    current_title = ''
    current_lines = []

    def _flush_chapter():
        nonlocal current_title, current_lines
        text = '\n'.join(current_lines)
        # 如果内容为空但有标题，也保留（可能是占位章）
        if current_title or text.strip():
            chapters.append({
                'title': current_title or '',
                'content': text,
            })
        current_title = ''
        current_lines = []

    for para in doc.paragraphs:
        text = para.text

        # 跳过空段落
        if not text.strip():
            current_lines.append('')
            continue

        # 检测段落样式是否为标题
        style_name = (para.style.name if para.style else '').lower()
        is_heading = (
            style_name.startswith('heading') or
            style_name.startswith('heading 1') or
            style_name == '标题' or
            style_name == '标题 1' or
            para.style.name == 'Heading 1'
        )

        if is_heading:
            _flush_chapter()
            current_title = text.strip()
            continue

        # 检测段前是否有分页符
        has_page_break = False
        for run in para.runs:
            for br in run._element.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    has_page_break = True
        # 也在段落属性中检测
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            for br in pPr.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    has_page_break = True

        if has_page_break and current_lines:
            current_lines.append('\f')

        # 检测首行缩进
        indent_val = None
        if para.paragraph_format.first_line_indent:
            indent_val = para.paragraph_format.first_line_indent

        line = text.strip()
        if indent_val and indent_val > Pt(0):
            line = '　　' + line

        current_lines.append(line)

    _flush_chapter()

    # 如果没有检测到章节，整篇作为一章
    if not chapters:
        chapters.append({
            'title': '',
            'content': '\n'.join(current_lines),
        })

    return {
        'title': chapters[0]['title'] if chapters else '',
        'chapters': chapters,
    }


# ── TXT 导出辅助 ────────────────────────────────────────────────

PAGE_BREAK_TXT = '\n\n***\n\n'


def content_to_txt(content: str) -> str:
    """将编辑器内容转为 TXT 格式，\\f → ***"""
    return content.replace('\f', PAGE_BREAK_TXT)


def build_book_txt(book) -> str:
    """全书合并为 TXT 字符串"""
    parts = [f'{book.title}\n{"=" * len(book.title)}\n']
    for ch in book.chapters:
        parts.append(f'\n{ch.title}\n{"-" * len(ch.title)}\n')
        parts.append(content_to_txt(ch.content or ''))
        parts.append('')
    return '\n'.join(parts)


# ── TXT 导入辅助 ────────────────────────────────────────────────

# 匹配导出时产生的分页标记：\\n\\n***\\n\\n
_PAGE_BREAK_RE = re.compile(r'\n{1,2}\*{3,5}\n{1,2}')


def _restore_page_breaks(text: str) -> str:
    """将 TXT 中的 *** 分页标记还原为 \\f"""
    return _PAGE_BREAK_RE.sub('\f', text)


def parse_txt(text: str, split_by_blank_lines: bool = False):
    """解析 TXT 文本。

    参数：
      text: 原始文本
      split_by_blank_lines: True 时按连续空行拆分为多章

    返回：
      [{'title': str, 'content': str}]

    会自动将 *** 分节线还原为 \\f 分页符。
    """
    # 还原分页符（导出时 \\f → ***）
    text = _restore_page_breaks(text)

    if not split_by_blank_lines:
        return [{'title': '', 'content': text}]

    # 按连续 3 个以上空行拆分
    parts = re.split(r'\n{3,}', text)
    chapters = []
    for part in parts:
        part = part.strip()
        if not part:
            continue
        # 尝试提取第一行作为标题
        lines = part.split('\n')
        title = ''
        content_start = 0
        if lines and len(lines[0].strip()) <= 50:
            title = lines[0].strip()
            content_start = 1
        chapters.append({
            'title': title,
            'content': '\n'.join(lines[content_start:]).strip(),
        })
    return chapters


# ── EPUB 导入 ──────────────────────────────────────────────────────

class _HTMLStripper(HTMLParser):
    """从 HTML 中提取纯文本，保留段落结构。"""
    def __init__(self):
        super().__init__()
        self.text = []
        self._skip = False

    def handle_starttag(self, tag, attrs):
        if tag in ('br', 'p', 'div', 'li', 'tr', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            self.text.append('\n')
        if tag in ('script', 'style', 'head', 'title'):
            self._skip = True

    def handle_endtag(self, tag):
        if tag in ('p', 'div', 'li', 'tr', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            self.text.append('\n')
        if tag in ('script', 'style', 'head', 'title'):
            self._skip = False

    def handle_data(self, data):
        if not self._skip:
            self.text.append(data)

    def get_text(self):
        raw = ''.join(self.text)
        # 压缩多个连续换行为双换行（段落分隔）
        raw = re.sub(r'\n{3,}', '\n\n', raw)
        # 清理每行多余空白
        lines = [line.strip() for line in raw.split('\n')]
        return '\n'.join(lines)


def parse_epub(filepath_or_stream):
    """解析 EPUB 文件，提取文本和章节结构。

    参数：
      filepath_or_stream: str 路径 或 io.BytesIO 对象

    返回：
      {
        'title': str,
        'chapters': [{'title': str, 'content': str}, ...]
      }

    映射规则：
      - EPUB spine 中的每个 XHTML 文档 → 一章
      - 文档中的 <h1> 优先作为章节标题，其次是 <title>
      - HTML 标签被剥离，保留段落换行
    """
    try:
        from ebooklib import epub as epub_lib
    except ImportError:
        raise ImportError('请安装 ebooklib：pip install ebooklib')

    book = epub_lib.read_epub(filepath_or_stream)

    # 提取书名
    title = ''
    titles = book.get_metadata('DC', 'title')
    if titles:
        title = titles[0][0]

    chapters = []
    stripper = _HTMLStripper()

    # 遍历 spine（阅读顺序）
    for item_id in book.spine:
        try:
            item_id_str = item_id[0] if isinstance(item_id, tuple) else item_id
        except (IndexError, TypeError):
            continue

        item = book.get_item_with_id(item_id_str)
        if item is None:
            continue

        # 只处理 XHTML/HTML 文档
        content_type = item.get_type()
        if content_type not in (
            'application/xhtml+xml',
            'text/html',
            'application/xml',
            'text/xml',
        ):
            # ebooklib 有时返回数字类型码，也可以检查 media_type 属性
            media = getattr(item, 'media_type', '')
            if not media or 'html' not in media and 'xml' not in media:
                continue

        raw = item.get_content().decode('utf-8', errors='replace')

        # 提取标题：优先 <h1>，其次 <title>
        ch_title = ''
        h1_match = re.search(r'<h1[^>]*>(.*?)</h1>', raw, re.IGNORECASE | re.DOTALL)
        if h1_match:
            ch_title = re.sub(r'<[^>]+>', '', h1_match.group(1)).strip()
        if not ch_title:
            title_match = re.search(r'<title[^>]*>(.*?)</title>', raw, re.IGNORECASE | re.DOTALL)
            if title_match:
                ch_title = re.sub(r'<[^>]+>', '', title_match.group(1)).strip()

        # 去除 <body> 之外的内容
        body_match = re.search(r'<body[^>]*>(.*?)</body>', raw, re.IGNORECASE | re.DOTALL)
        if body_match:
            raw = body_match.group(1)

        # 提取纯文本
        stripper.text = []
        stripper._skip = False
        stripper.feed(raw)
        text = stripper.get_text()

        if not text.strip():
            continue

        chapters.append({
            'title': ch_title,
            'content': text.strip(),
        })

    # 如果没有检测到章节（spine 为空或无 HTML 项），尝试回退读取全部
    if not chapters:
        chapters.append({
            'title': '',
            'content': title or '',
        })

    return {
        'title': title,
        'chapters': chapters,
    }
